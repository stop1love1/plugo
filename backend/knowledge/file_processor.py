"""File processor — extract text from PDF, DOCX, CSV, TXT, MD files."""

import csv
import io

from logging_config import logger


def extract_text(content: bytes, filename: str) -> str:
    """Extract text content from a file based on its extension."""
    name = filename.lower()

    if name.endswith((".txt", ".md")):
        return content.decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        return _extract_pdf(content)

    if name.endswith(".docx"):
        return _extract_docx(content)

    if name.endswith(".csv"):
        return _extract_csv(content)

    raise ValueError(f"Unsupported file type: {filename}")


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as err:
        raise ValueError("pypdf is not installed. Run: pip install pypdf") from err

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as e:
        raise ValueError(f"Could not read PDF file: {e!s}") from e

    parts = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text:
                parts.append(text.strip())
        except Exception as e:
            logger.warning("Failed to extract text from PDF page", page=i, error=str(e))
            continue
    return "\n\n".join(parts)


def _extract_docx(content: bytes) -> str:
    try:
        import docx
    except ImportError as err:
        raise ValueError("python-docx is not installed. Run: pip install python-docx") from err

    doc = docx.Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


MAX_CSV_ROWS = 10000


def _extract_csv(content: bytes) -> str:
    text = content.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    rows = []
    for i, row in enumerate(reader):
        if i >= MAX_CSV_ROWS:
            rows.append(f"... (truncated at {MAX_CSV_ROWS} rows)")
            break
        rows.append(" | ".join(row))
    return "\n".join(rows)
